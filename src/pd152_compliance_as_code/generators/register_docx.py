from __future__ import annotations

from pathlib import Path
from typing import Iterable

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency for some environments
    Document = None  # type: ignore

from pd152_compliance_as_code.domain.models import ComplianceConfig, ProcessingActivity

HEADERS = [
    "ID процесса",
    "Категории субъектов",
    "Цели",
    "Категории ПДн",
    "Правовые основания",
    "Получатели",
    "Местонахождение БД",
    "Сроки хранения",
    "Меры защиты",
]


def _format_list(items: Iterable[str]) -> str:
    return ", ".join(items)


def generate_register(config: ComplianceConfig, output_path: str | Path, locale: str = "ru") -> Path:
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX generation")

    doc = Document()
    doc.add_heading("Реестр операций по обработке ПДн", level=1)

    table = doc.add_table(rows=1, cols=len(HEADERS))
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(HEADERS):
        hdr_cells[idx].text = header

    for activity in config.processing_activities:
        row = table.add_row().cells
        row[0].text = activity.id
        row[1].text = _format_list(activity.subjects)
        row[2].text = _format_list(activity.purposes)
        row[3].text = _format_list(activity.data_categories)
        row[4].text = _format_list(activity.legal_basis)
        row[5].text = _format_list(activity.recipients)
        row[6].text = _format_list(activity.storage_locations)
        row[7].text = activity.retention.description
        measures = activity.security_measures
        row[8].text = _format_list(measures.org + measures.tech + measures.legal)

    out = Path(output_path)
    doc.save(out)
    return out
