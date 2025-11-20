from __future__ import annotations

from pathlib import Path
from typing import Iterable, List

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore

from pd152_compliance_as_code.domain.models import ComplianceConfig, ProcessingActivity


CONSENT_INTRO = (
    "Настоящий шаблон согласия подготовлен автоматически. Он не заменяет юридическую проверку и "
    "должен быть адаптирован под конкретную ситуацию."
)


def _format_lines(items: Iterable[str]) -> str:
    return ", ".join(items)


def generate_consents(
    config: ComplianceConfig,
    processes: List[str] | None,
    output_path: str | Path,
    policy_url: str | None = None,
) -> Path:
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX generation")

    selected_ids = set(processes) if processes else {p.id for p in config.processing_activities}
    doc = Document()
    doc.add_heading("Шаблоны согласий на обработку ПДн", level=1)
    doc.add_paragraph(CONSENT_INTRO)

    for activity in config.processing_activities:
        if activity.id not in selected_ids:
            continue
        doc.add_heading(activity.name, level=2)
        doc.add_paragraph("ФИО субъекта: __________________________")
        doc.add_paragraph("Контакты субъекта: _____________________")
        if policy_url:
            doc.add_paragraph(f"Ссылка на Политику: {policy_url}")
        doc.add_paragraph(f"Цели обработки: {_format_lines(activity.purposes)}")
        doc.add_paragraph(f"Категории ПДн: {_format_lines(activity.data_categories)}")
        doc.add_paragraph(f"Правовые основания: {_format_lines(activity.legal_basis)}")
        doc.add_paragraph("Согласие может быть отозвано по запросу субъекта.\n" "Проверьте текст с юристом перед использованием.")

    out = Path(output_path)
    doc.save(out)
    return out
