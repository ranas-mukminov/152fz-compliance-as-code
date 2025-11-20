from __future__ import annotations

from pathlib import Path
from typing import List

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore

from pd152_compliance_as_code.domain.models import ComplianceConfig, ProcessingActivity


DPA_NOTICE = (
    "Черновик поручения обработчику (DPA). Текст необходимо адаптировать и согласовать со "
    "сторонами. Пример не содержит цитат закона и служит основой для обсуждения."
)


def generate_dpa(config: ComplianceConfig, processors: List[str] | None, output_path: str | Path) -> Path:
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX generation")

    doc = Document()
    doc.add_heading("Черновик DPA / поручения обработчику", level=1)
    doc.add_paragraph(DPA_NOTICE)

    for processor in config.processors:
        if processors and processor.name not in processors:
            continue
        doc.add_heading(processor.name, level=2)
        doc.add_paragraph(f"Роль: {processor.role or 'обработчик'}")
        scope = _collect_scope(config.processing_activities)
        doc.add_paragraph(f"Предмет: обработка ПДн согласно процессам: {', '.join(scope)}")
        doc.add_paragraph(
            "Обязанности обработчика: действовать по поручению оператора, обеспечивать сохранность "
            "ПДн, использовать адекватные орг/тех/правовые меры, уведомлять об инцидентах."
        )
        doc.add_paragraph(
            "Проверка мер: оператор оценивает меры безопасности перед передачей и периодически "
            "пересматривает их актуальность."
        )

    out = Path(output_path)
    doc.save(out)
    return out


def _collect_scope(activities: List[ProcessingActivity]) -> List[str]:
    return [f"{a.id} ({a.name})" for a in activities]
